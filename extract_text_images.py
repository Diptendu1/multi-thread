"""
Extract text and images from .pdf, .docx/.doc, and .txt files.

Dependencies:
    pip install PyMuPDF python-docx docx2txt

    # legacy .doc (binary Word 97-2003) additionally needs LibreOffice on PATH
    # (used to convert .doc -> .docx transparently). On Debian/Ubuntu:
    #     sudo apt-get install libreoffice

Usage (CLI):
    python extract_text_images.py file1.pdf file2.docx notes.txt -o extracted/

Usage (library):
    from extract_text_images import extract
    result = extract("report.pdf", image_dir="extracted/report_images")
    print(result.text)
    print(result.images)   # list of saved image paths
"""
from __future__ import annotations

import argparse
import os
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass, field


@dataclass
class ExtractResult:
    source: str
    file_type: str
    text: str = ""
    images: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


# --------------------------------------------------------------------------- #
# Helpers
# --------------------------------------------------------------------------- #
def _ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path


def _base_name(path: str) -> str:
    return os.path.splitext(os.path.basename(path))[0]


# --------------------------------------------------------------------------- #
# PDF  (text + embedded raster images via PyMuPDF)
# --------------------------------------------------------------------------- #
def extract_pdf(path: str, image_dir: str) -> ExtractResult:
    import fitz  # PyMuPDF

    res = ExtractResult(source=path, file_type="pdf")
    _ensure_dir(image_dir)
    doc = fitz.open(path)
    text_parts: list[str] = []
    seen_xrefs: set[int] = set()

    for page_index in range(len(doc)):
        page = doc[page_index]
        text_parts.append(page.get_text("text"))

        for img in page.get_images(full=True):
            xref = img[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                pix = fitz.Pixmap(doc, xref)
                # Convert CMYK / alpha to RGB so it saves as PNG cleanly.
                if pix.n - pix.alpha >= 4:
                    pix = fitz.Pixmap(fitz.csRGB, pix)
                out = os.path.join(image_dir, f"{_base_name(path)}_p{page_index + 1}_x{xref}.png")
                pix.save(out)
                res.images.append(out)
                pix = None
            except Exception as exc:  # keep going on a single bad image
                res.warnings.append(f"page {page_index + 1} xref {xref}: {exc}")

    res.text = "\n".join(text_parts).strip()
    doc.close()
    return res


# --------------------------------------------------------------------------- #
# DOCX  (text via python-docx; images from the docx zip archive)
# --------------------------------------------------------------------------- #
def extract_docx(path: str, image_dir: str) -> ExtractResult:
    import docx  # python-docx

    res = ExtractResult(source=path, file_type="docx")
    _ensure_dir(image_dir)

    document = docx.Document(path)

    # Paragraph text.
    parts = [p.text for p in document.paragraphs]
    # Table text.
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    res.text = "\n".join(parts).strip()

    # Images live under word/media/ inside the docx (a zip archive).
    with zipfile.ZipFile(path) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
        for name in media:
            data = zf.read(name)
            out = os.path.join(image_dir, f"{_base_name(path)}_{os.path.basename(name)}")
            with open(out, "wb") as fh:
                fh.write(data)
            res.images.append(out)

    return res


# --------------------------------------------------------------------------- #
# DOC  (legacy binary Word) -> convert to .docx with LibreOffice, then reuse
# --------------------------------------------------------------------------- #
def extract_doc(path: str, image_dir: str) -> ExtractResult:
    res = ExtractResult(source=path, file_type="doc")

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        res.warnings.append(
            "Legacy .doc requires LibreOffice (soffice) on PATH to convert to .docx. "
            "Install it, or convert the file to .docx manually."
        )
        # Best-effort text-only fallback via docx2txt is not possible for .doc;
        # try `antiword` if present.
        antiword = shutil.which("antiword")
        if antiword:
            try:
                res.text = subprocess.check_output([antiword, path], text=True).strip()
            except Exception as exc:
                res.warnings.append(f"antiword failed: {exc}")
        return res

    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, path],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        converted = os.path.join(tmp, _base_name(path) + ".docx")
        if not os.path.exists(converted):
            res.warnings.append("LibreOffice conversion did not produce a .docx.")
            return res
        docx_res = extract_docx(converted, image_dir)

    # Merge, keeping original source/type.
    res.text = docx_res.text
    res.images = docx_res.images
    res.warnings.extend(docx_res.warnings)
    return res


# --------------------------------------------------------------------------- #
# TXT  (text only; no images)
# --------------------------------------------------------------------------- #
def extract_txt(path: str, image_dir: str) -> ExtractResult:
    res = ExtractResult(source=path, file_type="txt")
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        res.text = fh.read().strip()
    res.warnings.append("Text files contain no embedded images.")
    return res


# --------------------------------------------------------------------------- #
# Dispatcher
# --------------------------------------------------------------------------- #
_HANDLERS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_doc,
    ".txt": extract_txt,
}


def extract(path: str, image_dir: str | None = None) -> ExtractResult:
    """Extract text and images from a single file, dispatching on extension."""
    if not os.path.isfile(path):
        raise FileNotFoundError(path)
    ext = os.path.splitext(path)[1].lower()
    handler = _HANDLERS.get(ext)
    if handler is None:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {', '.join(_HANDLERS)}")
    if image_dir is None:
        image_dir = os.path.join(os.getcwd(), f"{_base_name(path)}_images")
    return handler(path, image_dir)


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #
def main(argv: list[str] | None = None) -> int:
    p = argparse.ArgumentParser(description="Extract text and images from PDF / DOC(X) / TXT files.")
    p.add_argument("files", nargs="+", help="Input files (.pdf, .docx, .doc, .txt).")
    p.add_argument("-o", "--outdir", default="extracted", help="Base output directory.")
    p.add_argument("--no-text-dump", action="store_true", help="Do not write extracted text to .txt sidecar files.")
    args = p.parse_args(argv)

    base_out = _ensure_dir(args.outdir)
    exit_code = 0

    for f in args.files:
        try:
            img_dir = os.path.join(base_out, f"{_base_name(f)}_images")
            result = extract(f, image_dir=img_dir)
        except Exception as exc:
            print(f"[ERROR] {f}: {exc}", file=sys.stderr)
            exit_code = 1
            continue

        if not args.no_text_dump and result.text:
            txt_path = os.path.join(base_out, f"{_base_name(f)}.txt")
            with open(txt_path, "w", encoding="utf-8") as fh:
                fh.write(result.text)

        print(f"{f}  [{result.file_type}]")
        print(f"  text: {len(result.text)} chars")
        print(f"  images: {len(result.images)} saved to {img_dir}")
        for w in result.warnings:
            print(f"  warning: {w}")

    return exit_code


if __name__ == "__main__":
    sys.exit(main())
